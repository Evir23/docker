import openpyxl
from docx import Document
import os
from datetime import datetime
from dateutil.relativedelta import relativedelta
from openpyxl.styles import Border, Side, Font
from xlsx2html import xlsx2html
import pdfkit


# Функция для чтения данных из протокола
def read_protocol(file_path):
    print("Начали")
    doc = Document(file_path)
    people_data = []
    chairman_name = None

    first_paragraph = doc.paragraphs[0].text

    # Предположим, что строка выглядит так: "Протокол № 7 от 20.09.2024"
    words = first_paragraph.split()

    # Ищем слово "№" и дату "от"
    protocol_number = words[words.index('№') + 1]  # Номер протокола
    protocol_date = words[words.index('от') + 1]  # Дата протокола

    # Проходим по всем абзацам в документе
    for paragraph in doc.paragraphs:
        if "Председатель комиссии:" in paragraph.text:
            chairman_name = paragraph.text.split("Председатель комиссии:")[1].strip()
            break

    # Предположим, что данные в протоколе находятся в таблице
    for table in doc.tables:
        for row in table.rows[1:]:
            cells = row.cells
            if len(cells) >= 7:
                person = {
                    'ФИО': cells[1].text.strip(),
                    'Профессия': cells[2].text.strip(),
                    'Компания': cells[3].text.strip(),
                    'Дата проверки': cells[5].text.strip(),
                    'Регистрационный номер': cells[6].text.strip(),
                    'Номер протокола': protocol_number,
                    'Дата протокола': protocol_date,
                    'Директор': chairman_name
                }
                people_data.append(person)

    return people_data


# Функция для создания Excel файла
def xlx_file(wb, ws, person, output_directory, company_name, group_number):
    ws['C4'] = f"{person['ФИО']}"
    font_12pt = Font(size=16)
    ws['C4'].font = font_12pt

    ws['C7'] = f"{person['Профессия']}"
    ws['C8'] = f"{company_name}"
    ws['B11'] = f"{person['Дата протокола']}"
    ws['G9'] = f"Основание: протокол № {person['Номер протокола']} от {person['Дата протокола']}"
    ws['G10'] = f"Директор {company_name}  ________________________"
    ws['G4'] = f"-присвоена {group_number} группа по безопасности работ на высоте."

    date_obj = datetime.strptime(person['Дата проверки'], "%d.%m.%Y")
    new_date_obj = date_obj + relativedelta(years=5)
    new_date_string = new_date_obj.strftime("%d.%m.%Y")
    ws['E11'] = new_date_string

    file_name = f"Удостоверение_{person['ФИО'].replace(' ', '_')}_{person['Регистрационный номер']}.xlsx"
    file_path = os.path.join(output_directory, file_name)

    wb.save(file_path)
    return file_path


# Удаление границ ячеек
def remove_cell_borders(ws):
    thin = Side(border_style=None, color='FFFFFF')
    no_border = Border(left=thin, right=thin, top=thin, bottom=thin)
    thick = Side(border_style="thick", color="000000")
    thick_border = Border(left=thick, right=thick, top=thick, bottom=thick)

    for row in ws.iter_rows():
        for cell in row:
            cell.border = no_border

    ws['A2'].border = thick_border
    # Добавляем толстую рамку вокруг диапазона A1:K19
    range_start = 'A1'
    range_end = 'J12'

    # Получаем координаты
    min_col, min_row, max_col, max_row = ws[range_start].column, ws[range_start].row, ws[range_end].column, ws[
        range_end].row

    # Применяем толстые границы только к краям диапазона
    for row in range(min_row, max_row + 1):
        for col in range(min_col, max_col + 1):
            cell = ws.cell(row=row, column=col)

            # Верхняя и нижняя границы
            if row == min_row:
                cell.border = Border(top=thick, left=cell.border.left, right=cell.border.right,
                                     bottom=cell.border.bottom)
            elif row == max_row:
                cell.border = Border(bottom=thick, left=cell.border.left, right=cell.border.right, top=cell.border.top)

            # Левая и правая границы
            if col == min_col:
                cell.border = Border(left=thick, top=cell.border.top, right=cell.border.right,
                                     bottom=cell.border.bottom)
            elif col == max_col:
                cell.border = Border(right=thick, top=cell.border.top, left=cell.border.left, bottom=cell.border.bottom)


# Конвертация Excel в PDF
def excel_to_pdf_with_no_borders(excel_file, output_pdf):
    wb = openpyxl.load_workbook(excel_file)
    ws = wb.active
    remove_cell_borders(ws)

    temp_excel_file = 'temp_no_borders.xlsx'
    wb.save(temp_excel_file)

    temp_html_file = 'temp_output.html'
    with open(temp_html_file, 'w', encoding='utf-8') as f:
        xlsx2html(temp_excel_file, f)

    path_to_wkhtmltopdf = r'C:\Program Files\wkhtmltopdf\bin\wkhtmltopdf.exe'
    config = pdfkit.configuration(wkhtmltopdf=path_to_wkhtmltopdf)

    pdfkit.from_file(temp_html_file, output_pdf, configuration=config)

    os.remove(temp_excel_file)
    os.remove(temp_html_file)

    print(f"PDF файл '{output_pdf}' успешно создан.")


# Основная функция для обработки протокола и генерации удостоверений
def process_documents(company_name, template, group_number, protocol_path, output_directory):
    print('Запуск обработки документов')
    try:
        people_data = read_protocol(protocol_path)
        temp_files = []  # Список для хранения временных файлов, которые нужно удалить

        for person in people_data:
            # Генерация Excel файла
            wb = openpyxl.load_workbook(template)
            ws = wb.active

            excel_file_path = xlx_file(wb, ws, person, output_directory, company_name, group_number)
            temp_files.append(excel_file_path)  # Добавляем путь к Excel файлу для последующего удаления

            # Генерация PDF файла
            pdf_file_path = os.path.join(output_directory, f"{person['ФИО'].replace(' ', '_')}.pdf")
            excel_to_pdf_with_no_borders(excel_file_path, pdf_file_path)

        # Очистка временных файлов после создания всех PDF
        for temp_file in temp_files:
            if os.path.exists(temp_file):
                os.remove(temp_file)  # Удаляем временные файлы
        print("Все временные файлы удалены.")

    except Exception as e:
        print(f"Ошибка в процессе обработки: {e}")
        raise
